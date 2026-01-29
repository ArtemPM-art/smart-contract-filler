# -*- coding: utf-8 -*-
"""
Smart Fill v3.1 Desktop Version
Based on Colab v3.1 Logic + V2 Desktop GUI
"""
import logging
import datetime
import tkinter.ttk as ttk # Нужно для выпадающего списка 
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os
import re
import sys
import tkinter as tk
from tkinter import filedialog, messagebox
GLOBAL_CONFIG = {
    'mode': 'std',
    'currency_symbol': '',
    'match_font': False,
    'log_dir': os.getcwd() # По умолчанию - папка скрипта
}
# --- Проверка библиотек ---
N2W_OK = False
try:
    from num2words import num2words as n2w_tool
    N2W_OK = True
except ImportError:
    logging.info("ВНИМАНИЕ: Библиотека num2words не найдена.")
    # Заглушка
    def n2w_tool(num, lang='ru'):
        return "ОШИБКА_БИБЛИОТЕКИ"

# Импортируем messagebox явно, чтобы использовать его в блоках except
from tkinter import messagebox

# --- НАСТРОЙКИ ШРИФТОВ ---
FONT_NAME = 'Times New Roman'
FONT_SIZE = 10

# Глобальная переменная для режима (как в V2)
SELECTED_MODE = 'std'

COL_MARKERS = {
    'qty': ['колво', 'количество', 'кол-во', 'к-во', 'кол.', 'ед.изм', 'шт', 'кол-', 'кол'],
    'price': ['ценабез', 'цена', 'цена(руб)', 'цена,руб', 'ценаза'],
    'sum_no_nds': ['суммабез', 'стоимостьбез', 'сумма', 'стоимость', 'всего'],
    'sum_nds_20': ['ндс20%', 'суммасндс20%', 'ндс20'],
    'sum_nds_22': ['ндс22%', 'суммасндс22%', 'ндс22', '22%'],
}

# =================================================================================================
# 1. GUI ИНТЕРФЕЙС 
# =================================================================================================

def select_options_gui(root):
    """
    Единое окно настроек: Режим, Форматирование, Логи.
    """
    # Результат, который вернет функция
    result = {'mode': None, 'currency': '', 'match_font': False}

    root.title("Smart Fill - Настройки")
    # --- НОВОЕ: УСТАНОВКА ИКОНКИ ---
    try:
        root.iconbitmap("icon.ico")
    except Exception:
        pass # Если иконки нет, просто будет стандартная синяя лапка Python

    # Размеры окна
    window_width = 400
    window_height = 450
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x_c = int((screen_width/2) - (window_width/2))
    y_c = int((screen_height/2) - (window_height/2))
    root.geometry(f"{window_width}x{window_height}+{x_c}+{y_c}")

    # --- 1. РЕЖИМ РАБОТЫ ---
    lbl_mode = tk.Label(root, text="1. Режим расчета:", font=("Arial", 10, "bold"))
    lbl_mode.pack(pady=(10, 5))
    
    mode_var = tk.StringVar(value="std")
    frame_modes = tk.Frame(root)
    frame_modes.pack(pady=5)
    
    # Радиокнопки
    rb_std = tk.Radiobutton(frame_modes, text="Стандартный", variable=mode_var, value="std")
    rb_std.pack(side=tk.LEFT, padx=10)
    rb_im = tk.Radiobutton(frame_modes, text="Проект ИМ", variable=mode_var, value="im")
    rb_im.pack(side=tk.LEFT, padx=10)

    # --- 2. ФОРМАТИРОВАНИЕ ---
    lbl_fmt = tk.Label(root, text="2. Форматирование:", font=("Arial", 10, "bold"))
    lbl_fmt.pack(pady=(15, 5))

    frame_fmt = tk.LabelFrame(root, text="Вид таблицы")
    frame_fmt.pack(padx=20, pady=5, fill="x")

    # Валюта
    tk.Label(frame_fmt, text="Символ валюты:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
    curr_combo = ttk.Combobox(frame_fmt, values=["Нет", "₽", "руб.", "руб"], state="readonly", width=10)
    curr_combo.current(1) # По умолчанию "₽"
    curr_combo.grid(row=0, column=1, padx=5, pady=5)

    # --- 3. ЛОГИРОВАНИЕ ---
    lbl_log = tk.Label(root, text="3. Где сохранить лог?", font=("Arial", 10, "bold"))
    lbl_log.pack(pady=(15, 5))
    
    # Переменная для пути (по умолчанию текущая папка)
    path_var = tk.StringVar(value=os.getcwd())
    
    def choose_dir():
        d = filedialog.askdirectory()
        if d:
            path_var.set(d)
            # Сразу обновляем конфиг
            GLOBAL_CONFIG['log_dir'] = d

    frame_log = tk.Frame(root)
    frame_log.pack(fill="x", padx=20)
    
    btn_dir = tk.Button(frame_log, text="📂 Выбрать папку...", command=choose_dir)
    btn_dir.pack(side=tk.LEFT, padx=5)
    
    lbl_path = tk.Label(frame_log, textvariable=path_var, fg="gray", font=("Arial", 8))
    lbl_path.pack(side=tk.LEFT, padx=5)

    # --- КНОПКА ЗАПУСКА ---
    def on_start():
        result['mode'] = mode_var.get()
        c_val = curr_combo.get()
        result['currency'] = "" if c_val == "Нет" else c_val
        root.quit() # Выходим из mainloop

    btn_start = tk.Button(root, text="ПРИМЕНИТЬ И ВЫБРАТЬ ФАЙЛ", command=on_start, 
                          bg="#4CAF50", fg="white", font=("Arial", 10, "bold"), height=2)
    btn_start.pack(pady=20, fill="x", padx=20)
    
    # Обработка закрытия крестиком
    def on_close():
        result['mode'] = None
        root.quit()
        
    root.protocol("WM_DELETE_WINDOW", on_close)
    
    # Запуск
    root.deiconify()
    root.mainloop()
    
    return result

# =================================================================================================
# 2. ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ (Из Colab v3)
# =================================================================================================
def setup_logger(folder_path):
    """
    Настраивает запись логов в файл.
    """
    log_filename = f"smart_fill_log_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.txt"
    full_path = os.path.join(folder_path, log_filename)
    
    # Настраиваем логгер: уровень INFO, запись в файл, кодировка utf-8
    # format определяет вид строки: "ВРЕМЯ - УРОВЕНЬ - СООБЩЕНИЕ"
    logging.basicConfig(
        filename=full_path,
        filemode='w',
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        encoding='utf-8',
        force=True # Важно для перезапуска конфигурации
    )
    
    # Также добавим вывод в консоль (пока тестируете), 
    # в EXE консоли не будет, но это не помешает
    console = logging.StreamHandler()
    console.setLevel(logging.INFO)
    formatter = logging.Formatter('%(message)s') # В консоль пишем просто текст
    console.setFormatter(formatter)
    logging.getLogger('').addHandler(console)
    
    logging.info(f"=== ЗАПУСК СКРИПТА: {datetime.datetime.now()} ===")
    logging.info(f"Лог-файл сохранен в: {full_path}")
    
def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent), 'paragraph'
        elif child.tag.endswith('tbl'):
            yield docx.table.Table(child, parent), 'table'

def clean_header(text):
    if not text: return ""
    text = text.lower()
    # Замена латиницы на кириллицу (фикс для 'C'ены)
    replacements = {
        'a': 'а', 'c': 'с', 'e': 'е', 'o': 'о', 'p': 'р',
        'x': 'х', 'y': 'у', 'k': 'к', 'h': 'н', 'b': 'в', 'm': 'м'
    }
    for lat, cyr in replacements.items():
        text = text.replace(lat, cyr)
    text = text.replace('\n', '').replace('\r', '').replace('\v', '').replace('\t', '')
    cleaned = re.sub(r'[\s\-\u00AD\.\,\:\(\)]', '', text)
    return cleaned

def clean_number(s):
    if not isinstance(s, str): return 0.0
    s = re.sub(r'[₽рРrRубa-zA-Zа-яА-Я]', '', s)
    s = re.sub(r'\s+', '', s)
    s = s.replace(',', '.')
    s = re.sub(r'[^\d.]', '', s)
    if not s: return 0.0
    if s.count('.') > 1:
        parts = s.split('.')
        s = "".join(parts[:-1]) + '.' + parts[-1]
    try: return float(s)
    except: return 0.0

def format_money_full(val):
    return "{:,.2f}".format(val).replace(",", " ").replace(".", ",")

def format_money_int(val):
    return "{:,.0f}".format(val).replace(",", " ")

def set_cell(cell, text):
    try:
        cell.text = ""
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(text))
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except Exception as e:
        logging.info(f"Ошибка записи в ячейку: {e}")

def clean_old_summaries(doc):
    logging.info("Очистка старых итогов...")
    keywords = [
        "Итого сумма составляет",
        "Итого сумма за Пусконаладочные",
        "Итого за программное обеспечение",
        "Общая цена настоящей спецификации",
    ]
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        for k in keywords:
            if k in txt:
                paragraphs_to_delete.append(p)
                break
    for p in paragraphs_to_delete:
        try: p._element.getparent().remove(p._element)
        except: pass

# Глобальная переменная, чтобы показать ошибку num2words только 1 раз за запуск
_n2w_error_shown = False 

def generate_text_parts(amount):
    global _n2w_error_shown
    rubles = int(amount)
    kopecks = int(round((amount - rubles) * 100))
    amount_str = format_money_int(rubles)
    
    text_sum = ""
    try:
        text_sum = n2w_tool(rubles, lang='ru').capitalize()
    except Exception as e:
        error_msg = f"Сбой num2words: {e}"
        logging.error(error_msg)
        text_sum = f"ОШИБКА_ПЕРЕВОДА"
        
        # Если это первая ошибка такого рода - покажем пользователю окно
        if not _n2w_error_shown:
            messagebox.showerror("Ошибка перевода числа", 
                                 f"Не удалось перевести число {rubles} в текст.\n\nОшибка: {e}\n\n"
                                 "Возможно, библиотека не установлена или число слишком большое.")
            _n2w_error_shown = True # Больше не показывать это окно в текущем сеансе
            
    return amount_str, text_sum, f"{kopecks:02d}"

# =================================================================================================
# 3. УПРАВЛЕНИЕ ПРОБЕЛАМИ (Из Colab v3 - New Feature)
# =================================================================================================

def is_xml_paragraph_empty(element):
    """
    Усиленная проверка: считает параграф пустым, даже если там
    есть неразрывные пробелы (\xa0) или просто пробелы.
    """
    if element is None: return False
    if not element.tag.endswith('p'): return False
    text = ""
    for node in element.itertext():
        text += node
    clean_text = text.replace('\xa0', ' ').strip()
    return len(clean_text) == 0

def add_empty_p_xml(parent):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    r.append(t)
    p.append(r)
    return p

def manage_spacing_before_element(item):
    """Обеспечивает ровно 1 пробел ПЕРЕД элементом."""
    element = item._element if hasattr(item, '_element') else item
    parent = element.getparent()
    prev = element.getprevious()
    empty_paragraphs_above = []

    while prev is not None:
        if is_xml_paragraph_empty(prev):
            empty_paragraphs_above.append(prev)
            prev = prev.getprevious()
        else:
            break

    count = len(empty_paragraphs_above)
    if count == 0:
        new_p = add_empty_p_xml(parent)
        element.addprevious(new_p)
    elif count > 1:
        for p_to_remove in empty_paragraphs_above[1:]:
            parent.remove(p_to_remove)

def manage_spacing_after_block(last_element):
    """Обеспечивает ровно 1 пробел ПОСЛЕ элемента."""
    if last_element is None: return
    element = last_element._element if hasattr(last_element, '_element') else last_element
    parent = element.getparent()

    next_el = element.getnext()
    empty_paragraphs_below = []

    while next_el is not None:
        if is_xml_paragraph_empty(next_el):
            empty_paragraphs_below.append(next_el)
            next_el = next_el.getnext()
        else:
            break

    count = len(empty_paragraphs_below)

    if count == 0:
        new_p = add_empty_p_xml(parent)
        element.addnext(new_p)
        # logging.info("    [SPACING] Добавлен пропущенный пробел ПОСЛЕ блока.") # Debug off

    elif count > 1:
        # Оставляем только первый найденный (ближайший), остальные удаляем
        for p_to_remove in empty_paragraphs_below[1:]:
            parent.remove(p_to_remove)
        # logging.info(f"    [SPACING] Удалено {count - 1} лишних пробелов ПОСЛЕ блока.") # Debug off

def insert_clean_paragraph_after(sibling_xml, text, parent_obj):
    """Вставляет параграф с текстом сразу после указанного XML-элемента без лишних отступов."""
    new_p_xml = OxmlElement("w:p")
    sibling_xml.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, parent_obj)

    run = new_p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    return new_p

# =================================================================================================
# 4. БИЗНЕС-ЛОГИКА (Из Colab v3)
# =================================================================================================

def get_table_text_preview(table):
    raw_text = ""
    limit = min(15, len(table.rows))
    for r in table.rows[:limit]:
        for c in r.cells:
            raw_text += clean_header(c.text)
    return raw_text, raw_text[:70]

def determine_table_type(raw_text_clean):
    if 'пусконалад' in raw_text_clean or 'пнр' in raw_text_clean or 'монтаж' in raw_text_clean:
        return 'pnr'

    soft_keywords = [
        'лицензия', 'программное', 'подписка',
        'неконкурентная', 'экземпляр', 'активаци', 'MDT', 'Промышленный интернет вещей'
    ]
    if any(word in raw_text_clean for word in soft_keywords):
        return 'soft'

    return 'equip'

def process_table(table, table_index):
    if not table.rows: return None

    # 1. Читаем текст
    raw_text, preview = get_table_text_preview(table)

    # 2. Фильтр пустых таблиц
    if len(raw_text) < 5:
        return None

    t_type = determine_table_type(raw_text)
    logging.info(f"\n--- Обработка таблицы №{table_index + 1} (Тип: {t_type.upper()}) ---")

    header_row_idx = -1
    idxs = {}

    # Ищем заголовки
    search_depth = min(15, len(table.rows))
    for r_idx in range(search_depth):
        headers = [clean_header(c.text) for c in table.rows[r_idx].cells]
        temp_idxs = {}
        for key, markers in COL_MARKERS.items():
            for i, h in enumerate(headers):
                if any(m in h for m in markers) and key not in temp_idxs:
                    temp_idxs[key] = i

        if ('qty' in temp_idxs) and ('price' in temp_idxs or 'sum_nds_22' in temp_idxs):
            idxs = temp_idxs
            header_row_idx = r_idx
            logging.info(f"    Заголовок найден в строке {r_idx}. Столбцы: {idxs}")
            break

    if not idxs:
        logging.info(f"    ВНИМАНИЕ: Не удалось найти заголовки. Пропускаем.")
        return None

    has_nds_22 = 'sum_nds_22' in idxs
    has_nds_20 = 'sum_nds_20' in idxs
    total_no_nds = 0.0

    last_row = len(table.rows)
    for i in range(header_row_idx + 1, len(table.rows)):
        txt = "".join([c.text.lower() for c in table.rows[i].cells[:3]])
        if "итого" in txt or "всего" in txt:
            last_row = i
            break

    # --- ПОДГОТОВКА ВАЛЮТЫ ИЗ НАСТРОЕК ---
    user_curr = GLOBAL_CONFIG.get('currency_symbol', '')
    curr_suffix = f" {user_curr}" if user_curr else ""
    logging.info(f"    Используемый символ валюты: '{user_curr}'")

    processed_ids = set()

    # --- ЦИКЛ ПО СТРОКАМ ---
    for i in range(header_row_idx + 1, last_row):
        row = table.rows[i]
        if len(row.cells) <= max(idxs.values()): continue

        cell_qty = row.cells[idxs['qty']]
        if 'price' in idxs:
            cell_price = row.cells[idxs['price']]
        else:
            cell_price = row.cells[idxs['qty']]

        if cell_price._tc in processed_ids: continue
        processed_ids.add(cell_price._tc)

        # ЧИТАЕМ ЗНАЧЕНИЯ
        qty = clean_number(cell_qty.text)
        price_val = 0.0
        if 'price' in idxs:
            price_val = clean_number(cell_price.text)

        val_22 = 0.0
        if has_nds_22:
            val_22 = clean_number(row.cells[idxs['sum_nds_22']].text)

        # !!! ЛОГИРОВАНИЕ СЧИТАННЫХ ДАННЫХ !!!
        # Берем название товара (обычно 2-я колонка, индекс 1), чтобы в логе было понятно
        try:
            item_name = row.cells[1].text.strip().replace('\n', ' ')[:30] + "..."
        except:
            item_name = "Товар..."
            
        logging.info(f"    [Стр {i}] '{item_name}' -> ВИЖУ: Кол={qty}, Цена={price_val}, НДС22={val_22}")

        if price_val <= 0.0001 and val_22 <= 0.0001:
            # Если цена 0, пишем в лог, что пропускаем
            if qty > 0:
                logging.info(f"        -> Пропуск (нет цены и нет НДС)")
            continue

        # ПНР Фикс
        is_header_row = False
        try:
            if re.search(r'[а-яА-Яa-zA-Z]{3,}', cell_qty.text):
                is_header_row = True
        except: pass

        if not is_header_row and t_type == 'pnr':
            if qty <= 0.01:
                qty = 1.0
                set_cell(cell_qty, "1")
                logging.info(f"        -> ПНР: Авто-установка кол-ва = 1")

        curr_sum_no_nds = 0.0
        calculated = False

        # Расчет от НДС (обратный)
        if has_nds_22 and val_22 > 0.01:
            curr_sum_no_nds = round(val_22 / 1.22, 2)
            if qty > 0 and 'price' in idxs:
                price_no_nds = round(curr_sum_no_nds / qty, 2)
                set_cell(cell_price, format_money_full(price_no_nds) + curr_suffix)
                logging.info(f"        -> Вычислена Цена (из НДС): {price_no_nds}")
            
            set_cell(row.cells[idxs['sum_nds_22']], format_money_full(val_22) + curr_suffix)
            calculated = True

        # Расчет прямой
        if not calculated and 'price' in idxs:
            curr_sum_no_nds = round(qty * price_val, 2)
            logging.info(f"        -> Расчет суммы: {qty} * {price_val} = {curr_sum_no_nds}")
            
            if has_nds_22:
                nds22 = round(curr_sum_no_nds * 1.22, 2)
                set_cell(row.cells[idxs['sum_nds_22']], format_money_full(nds22) + curr_suffix)
                logging.info(f"        -> Вычислен НДС22: {nds22}")

        if 'sum_no_nds' in idxs:
            set_cell(row.cells[idxs['sum_no_nds']], format_money_full(curr_sum_no_nds) + curr_suffix)
            
        if has_nds_20:
            nds20 = round(curr_sum_no_nds * 1.20, 2)
            set_cell(row.cells[idxs['sum_nds_20']], format_money_full(nds20) + curr_suffix)

        total_no_nds += curr_sum_no_nds

    # Итоги таблицы
    logging.info(f"    ИТОГ ТАБЛИЦЫ (без НДС): {total_no_nds}")
    
    if last_row < len(table.rows):
        cells = table.rows[last_row].cells
        try:
            if 'sum_no_nds' in idxs:
                set_cell(cells[idxs['sum_no_nds']], format_money_full(total_no_nds) + curr_suffix)
            if has_nds_20:
                t20 = round(total_no_nds * 1.20, 2)
                set_cell(cells[idxs['sum_nds_20']], format_money_full(t20) + curr_suffix)
            if has_nds_22:
                t22 = round(total_no_nds * 1.22, 2)
                set_cell(cells[idxs['sum_nds_22']], format_money_full(t22) + curr_suffix)
        except Exception as e:
            logging.info(f"    Ошибка итогов: {e}")

    return {'sum': total_no_nds, 'type': t_type, 'table': table}

def add_summary_after_table(table, amount, t_type):
    if amount <= 0.001: return

    s, t, k = generate_text_parts(amount)
    text = ""

    if t_type == 'equip':
        text = f"Итого сумма составляет {s} ({t}) рублей, {k} копеек, кроме того НДС в соответствии действующим законодательством РФ."
    elif t_type == 'pnr':
        text = f"Итого сумма за Пусконаладочные работы составляет {s} ({t}) руб., {k} копеек, кроме того НДС в соответствии с действующим законодательством РФ."
    elif t_type == 'soft':
        text = f"Итого за программное обеспечение сумма составляет (приводится справочно): {s} ({t}) руб, {k} копеек, НДС не облагается."

    if text:
        # 1. Вставляем текст
        new_p = insert_clean_paragraph_after(table._element, text, table._parent)

        # 2. ПРИНУДИТЕЛЬНО чистим всё вокруг этой новой строки
        manage_spacing_before_element(new_p)
        manage_spacing_after_block(new_p)
        logging.info(f"    -> Добавлен итог под таблицей.")

def finalize_section_totals(doc, totals, last_processed_table, mode):
    # 1. Считаем математику
    grand_total = totals['equip'] + totals['pnr']
    
    # В стандартном режиме добавляем софт, в ИМ - нет
    if mode == 'std': 
        grand_total += totals['soft']

    # Если сумма 0, выходим
    if grand_total <= 0.01: return

    logging.info(f"\n--- ИТОГ РАЗДЕЛА: Общая цена = {grand_total} (Режим: {mode}) ---")
    
    s, t, k = generate_text_parts(grand_total)
    grand_text = f"Общая цена настоящей спецификации составляет {s} ({t}) рублей, {k} копеек, кроме того НДС в соответствии действующим законодательством РФ."

    if not last_processed_table:
        return

    # --- ОПРЕДЕЛЕНИЕ МЕСТА ВСТАВКИ ---
    target_element = None
    insert_before = False # Флаг: вставлять ДО или ПОСЛЕ найденного элемента?

    # ЛОГИКА ДЛЯ РЕЖИМА "IM" (Ищем место ПЕРЕД софтом)
    if mode == 'im':
        # Начинаем от последней таблицы и идем ВВЕРХ
        current = last_processed_table._element
        
        # Ищем не более 50 элементов вверх (чтобы не уйти в другой раздел)
        for _ in range(50):
            prev = current.getprevious()
            if prev is None: break
            
            # Проверяем текст параграфа
            try:
                # Конвертируем XML в параграф для проверки текста
                if prev.tag.endswith('p'):
                    p_obj = Paragraph(prev, last_processed_table._parent)
                    text = p_obj.text.lower()
                    # Ключевая фраза из вашего скриншота
                    if "программное обеспечение" in text and "подписке" in text:
                        target_element = prev
                        insert_before = True
                        logging.info("    [IM] Найдено начало раздела ПО. Вставляем итог ПЕРЕД ним.")
                        break
                
                # Если наткнулись на другую таблицу - стоп, мы ушли слишком далеко
                if prev.tag.endswith('tbl'):
                    break
            except: pass
            current = prev

    # ЛОГИКА ДЛЯ РЕЖИМА "STD" (Или если в IM не нашли софт) - Вставляем в самый низ
    if target_element is None:
        insert_before = False
        current_element = last_processed_table._element
        target_element = current_element

        # Проходим вниз, чтобы найти последний элемент текущего блока (пропускаем "Итого...")
        for _ in range(10):
            next_sib = target_element.getnext()
            if next_sib is None or not next_sib.tag.endswith('p'): break

            try:
                p_obj = Paragraph(next_sib, last_processed_table._parent)
                text = p_obj.text.strip()
            except: text = ""

            keywords_intermediate = ["Итого за программное", "Итого сумма за", "Итого сумма составляет", "Итого:", "Всего:"]
            if not text or any(k in text for k in keywords_intermediate):
                target_element = next_sib
            else:
                break

    # --- ВСТАВКА ---
    if insert_before:
        # Вставка ПЕРЕД элементом (для IM)
        new_p_xml = OxmlElement("w:p")
        target_element.addprevious(new_p_xml)
        new_grand_p = Paragraph(new_p_xml, last_processed_table._parent)
        
        # Наполняем текстом
        run = new_grand_p.add_run(grand_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        
    else:
        # Вставка ПОСЛЕ элемента (для STD)
        new_grand_p = insert_clean_paragraph_after(target_element, grand_text, last_processed_table._parent)
        logging.info("    [STD] Вставка Общего итога в конец блока.")

    # --- ФИНАЛЬНОЕ УПРАВЛЕНИЕ ПРОБЕЛАМИ ---
    # Чистим отступы, чтобы было красиво
    manage_spacing_before_element(new_grand_p)
    manage_spacing_after_block(new_grand_p)

# =================================================================================================
# 5. MAIN (DESKTOP)
# =================================================================================================

def main():
    # --- 0. ПРОВЕРКА БИБЛИОТЕК GUI ---
    root = tk.Tk()
    
    # Если при импорте возникла проблема, сообщаем сразу
    if not N2W_OK:
        messagebox.showwarning("Внимание", 
                               "Библиотека 'num2words' не найдена!\n\n"
                               "Скрипт будет работать, но сумма прописью \n"
                               "будет заменена на слово 'ОШИБКА_БИБЛИОТЕКИ'.")
# --- НОВОЕ: УСТАНОВКА ИКОНКИ ДЛЯ ГЛАВНОГО ОКНА ---
    try:
        root.iconbitmap("icon.ico")
    except Exception:
        pass
# -------------------------------------------------
    
    # Вызываем новую функцию вместо старой select_mode_gui
    options = select_options_gui(root)
    
    # Если нажали крестик (mode is None)
    if not options['mode']:
        root.destroy()
        return

    # Сохраняем настройки
    GLOBAL_CONFIG['mode'] = options['mode']
    GLOBAL_CONFIG['currency_symbol'] = options['currency']
    GLOBAL_CONFIG['match_font'] = options['match_font']
    # GLOBAL_CONFIG['log_dir'] уже обновился внутри функции GUI
    
    root.withdraw()

    # --- 2. ВКЛЮЧАЕМ ЛОГИРОВАНИЕ ---
    setup_logger(GLOBAL_CONFIG['log_dir'])
    
    logging.info(f"Настройки приняты: {GLOBAL_CONFIG}")

    # --- 3. ВЫБОР ФАЙЛА ---
    root.update() 
    file_path = filedialog.askopenfilename(
        parent=root, 
        title="Выберите файл .docx",
        filetypes=[("Word Documents", "*.docx")]
    )

    if not file_path:
        logging.warning("Файл не выбран. Отмена.")
        root.destroy()
        return

    logging.info(f"Начинаем обработку файла: {file_path}")

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logging.critical(f"Критическая ошибка открытия файла: {e}")
        messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")
        root.destroy()
        return

    clean_old_summaries(doc)

    current_totals = {'equip': 0.0, 'pnr': 0.0, 'soft': 0.0}
    last_item_in_section = None
    table_global_index = 0

    all_items = list(iter_block_items(doc))

    # --- ОСНОВНОЙ ЦИКЛ ---
    for item, item_type in all_items:
        if item_type == 'paragraph':
            try:
                text = item.text.lower().strip()
            except: continue

            if "спецификация" in text and ("№" in text or "номер" in text):
                # !!! ИСПРАВЛЕНИЕ 1: Добавлен аргумент GLOBAL_CONFIG['mode']
                finalize_section_totals(doc, current_totals, last_item_in_section, GLOBAL_CONFIG['mode'])
                
                logging.info(f"--- Нашел раздел: {item.text[:40]}... Сброс итогов. ---")
                current_totals = {'equip': 0.0, 'pnr': 0.0, 'soft': 0.0}
                last_item_in_section = None

        elif item_type == 'table':
            manage_spacing_before_element(item)
            res = process_table(item, table_global_index)
            table_global_index += 1

            if res:
                add_summary_after_table(item, res['sum'], res['type'])
                current_totals[res['type']] += res['sum']
                last_item_in_section = item
                
                # Отступы
                next_el = item._element.getnext()
                target_element_for_spacing = item
                if next_el is not None and next_el.tag.endswith('p'):
                    if not is_xml_paragraph_empty(next_el):
                        target_element_for_spacing = next_el
                manage_spacing_after_block(target_element_for_spacing)

    # !!! ИСПРАВЛЕНИЕ 2: Добавлен аргумент GLOBAL_CONFIG['mode']
    finalize_section_totals(doc, current_totals, last_item_in_section, GLOBAL_CONFIG['mode'])

    # --- СОХРАНЕНИЕ ---
    dir_name, file_name = os.path.split(file_path)
    mode_str = GLOBAL_CONFIG['mode'].upper()
    new_name = file_name.replace(".docx", f"_ГОТОВЫЙ_{mode_str}.docx")
    full_save_path = os.path.join(dir_name, new_name)

    try:
        doc.save(full_save_path)
        logging.info(f"УСПЕХ! Файл сохранен: {full_save_path}")
        messagebox.showinfo("Готово", f"Готово!\nЛог сохранен в:\n{GLOBAL_CONFIG['log_dir']}")
    except Exception as e:
        logging.error(f"Ошибка сохранения: {e}")
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")

    root.destroy()

# =================================================================================================
# ЗАПУСК (С GUI-обработчиком ошибок)
# =================================================================================================

if __name__ == "__main__":
    try:
        # Проверка критически важных библиотек
        import docx
        import tkinter
        
        # Запуск
        main()
        
    except ImportError as e:
        # Эта ошибка сработает, если нет docx или tkinter
        # Мы пытаемся создать минимальное окно tkinter, чтобы показать ошибку
        try:
            import tkinter as tk
            from tkinter import messagebox
            root = tk.Tk()
            root.withdraw() # Скрываем основное окно
            messagebox.showerror("Критическая ошибка запуска", 
                                 f"Отсутствуют обязательные библиотеки!\n\nОшибка: {e}\n\n"
                                 "Убедитесь, что установлен python-docx.")
        except:
            # Если даже tkinter нет, то только консоль (но в exe этого не видно)
            print("CRITICAL ERROR: Libraries missing and cannot create GUI.")
            
    except Exception as e:
        # ЛОВИМ ЛЮБЫЕ ДРУГИЕ ОШИБКИ (Например, ошибки в логике кода)
        import traceback
        err_trace = traceback.format_exc()
        
        logging.critical(f"НЕОБРАБОТАННАЯ ОШИБКА: {e}")
        logging.critical(err_trace)
        
        # Пытаемся показать окно с ошибкой
        try:
            import tkinter as tk
            from tkinter import messagebox
            # Если root еще не создан или уничтожен, создаем новый для вывода ошибки
            if 'root' not in locals() or not tk._default_root:
                root = tk.Tk()
                root.withdraw()
                
            messagebox.showerror("Критическая ошибка программы", 
                                 f"Произошла неожиданная ошибка, программа будет закрыта.\n\n"
                                 f"Текст ошибки:\n{e}\n\n"
                                 f"Подробности записаны в лог-файл.")
        except:
            print(f"CRITICAL GUI FAIL: {e}")